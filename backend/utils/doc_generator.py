from docx import Document
from docx.shared import Inches
import zipfile
import os

def generate_all_docs(parsed, form, zip_path):
    """
    生成五份标准单据：提单、发票、箱单、受益人证明、汇票
    """
    # 提取数据
    shipper = form.get("shipper") or parsed["beneficiary"]
    consignee = form.get("consignee") or parsed["applicant"]
    notify = form.get("notify") or parsed["applicant"]
    vessel = form.get("vessel") or "VESSEL / VOY"
    goods = form.get("goods") or parsed["goods_description"]
    amount = parsed["amount"] or "USD 0.00"
    inv_no = form.get("invoice_no") or "INV-2026-001"
    lc_no = parsed["lc_number"] or "LC-XXXX"
    
    # 创建临时目录存放docx
    tmp_dir = "/tmp/lc_docs"
    os.makedirs(tmp_dir, exist_ok=True)
    
    def add_basic_header(doc, title):
        doc.add_heading(title, 0)
        doc.add_paragraph(f"L/C No.: {lc_no}")
        doc.add_paragraph(f"Invoice No.: {inv_no}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph("\n")
    
    # 1. 提单
    doc = Document()
    add_basic_header(doc, "BILL OF LADING (DRAFT)")
    doc.add_paragraph(f"Shipper: {shipper}")
    doc.add_paragraph(f"Consignee: TO ORDER OF ISSUING BANK / {consignee}")
    doc.add_paragraph(f"Notify Party: {notify}")
    doc.add_paragraph(f"Vessel: {vessel}")
    doc.add_paragraph(f"Port of Loading: {parsed.get('port_loading', '')}")
    doc.add_paragraph(f"Port of Discharge: {parsed.get('port_discharge', '')}")
    doc.add_paragraph(f"Goods: {goods}")
    doc.add_paragraph("\n_________________________")
    doc.add_paragraph("Signed by Carrier")
    doc.save(f"{tmp_dir}/Bill_of_Lading.docx")
    
    # 2. 发票
    doc = Document()
    add_basic_header(doc, "COMMERCIAL INVOICE")
    doc.add_paragraph(f"Beneficiary: {shipper}")
    doc.add_paragraph(f"Applicant: {consignee}")
    doc.add_paragraph(f"Goods: {goods}")
    doc.add_paragraph(f"Total Amount: {amount}")
    doc.add_paragraph("\n_________________________")
    doc.add_paragraph("Signed by Beneficiary")
    doc.save(f"{tmp_dir}/Invoice.docx")
    
    # 3. 装箱单
    doc = Document()
    add_basic_header(doc, "PACKING LIST")
    doc.add_paragraph(f"Marks: N/M")
    doc.add_paragraph(f"Description: {goods}")
    doc.add_paragraph(f"Qty: 120 PCS (待填)")
    doc.add_paragraph(f"G.W.: 2450 KG (待填)")
    doc.save(f"{tmp_dir}/Packing_List.docx")
    
    # 4. 受益人证明
    doc = Document()
    add_basic_header(doc, "BENEFICIARY CERTIFICATE")
    doc.add_paragraph("We hereby certify that one set of non-negotiable shipping documents has been sent to the applicant within 3 working days after shipment.")
    doc.add_paragraph(f"\n{shipper}")
    doc.add_paragraph("\n_________________________")
    doc.add_paragraph("Date: " + datetime.now().strftime('%Y-%m-%d'))
    doc.save(f"{tmp_dir}/Beneficiary_Certificate.docx")
    
    # 5. 汇票（两联）
    doc = Document()
    add_basic_header(doc, "BILL OF EXCHANGE (FIRST)")  # 简化，可单独做第二联
    # 完整汇票需包含金额大写，此处省略（可扩展）
    doc.save(f"{tmp_dir}/Bill_of_Exchange.docx")
    
    # 打包ZIP
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in os.listdir(tmp_dir):
            zipf.write(os.path.join(tmp_dir, file), file)
    
    # 清理
    for file in os.listdir(tmp_dir):
        os.remove(os.path.join(tmp_dir, file))
    os.rmdir(tmp_dir)
